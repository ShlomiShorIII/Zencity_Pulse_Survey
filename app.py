import streamlit as st
from supabase import create_client, Client
import pandas as pd
from docx import Document
import re

url = st.secrets["supabase"]["url"]
key = st.secrets["supabase"]["key"]
supabase: Client = create_client(url, key)

st.set_page_config(page_title="Survey Builder", layout="centered")
st.image("assets/zencity_logo.png", width=150)
st.title("📝 Pulse Survey Builder")

survey_title = st.text_input("Survey Title")
survey_intro = st.text_area("Survey Introduction")

@st.cache_data
def load_data():
    cats = supabase.table("categories").select("*").execute().data
    subs = supabase.table("subcategories").select("*").execute().data
    links = supabase.table("category_subcategory").select("*").execute().data
    return pd.DataFrame(cats), pd.DataFrame(subs), pd.DataFrame(links)

categories_df, subcategories_df, cat_sub_links_df = load_data()
cat_sub_links = cat_sub_links_df.to_dict(orient="records")

category_names = {row['category_name']: row['category_id'] for _, row in categories_df.iterrows()}
selected_cat_name = st.selectbox("Select a Category", list(category_names.keys()))
selected_cat_id = category_names[selected_cat_name]
linked_sub_ids = [link['subcategory_id'] for link in cat_sub_links if link['category_id'] == selected_cat_id]
filtered_subs = subcategories_df[subcategories_df['subcategory_id'].isin(linked_sub_ids)]

selected_sub_id = None
if not filtered_subs.empty:
    sub_names = {row['subcategory_name']: row['subcategory_id'] for _, row in filtered_subs.iterrows()}
    selected_sub_name = st.selectbox("Select a Subcategory", list(sub_names.keys()))
    selected_sub_id = sub_names[selected_sub_name]
else:
    st.warning("No subcategories found for the selected category.")

def load_questions(cat_id, sub_id):
    open_qs = supabase.table("open_questions").select("open_question_id, question_text")\
        .eq("category_id", cat_id).eq("subcategory_id", sub_id).execute().data
    closed_qs = supabase.table("closed_questions").select("closed_question_id, question_text")\
        .eq("category_id", cat_id).eq("subcategory_id", sub_id).execute().data
    return open_qs, closed_qs

if "db_questions" not in st.session_state:
    st.session_state["db_questions"] = []

if selected_sub_id:
    open_qs, closed_qs = load_questions(selected_cat_id, selected_sub_id)
    for q in open_qs:
        key = f"open_{q['open_question_id']}"
        if st.checkbox(f"[Open] {q['question_text']}", key=key):
            if key not in [q['id'] for q in st.session_state.db_questions]:
                st.session_state.db_questions.append({"id": key, "type": "Open", "text": q['question_text'], "options": []})

    for q in closed_qs:
        answers = supabase.table("closed_questions_answers").select("answer_option")\
            .eq("closed_question_id", q['closed_question_id']).execute().data
        options = [a['answer_option'] for a in answers]
        key = f"closed_{q['closed_question_id']}"
        if st.checkbox(f"[Closed] {q['question_text']}", key=key):
            if key not in [q['id'] for q in st.session_state.db_questions]:
                st.session_state.db_questions.append({"id": key, "type": "Closed", "text": q['question_text'], "options": options})

if "new_questions" not in st.session_state:
    st.session_state["new_questions"] = []

if st.button("➕ Add Custom Question"):
    st.session_state.new_questions.append({"type": "Open", "text": "", "options": []})
    st.rerun()

st.subheader("Edit Questions")
edited_questions = []

for i, q in enumerate(st.session_state.db_questions + st.session_state.new_questions):
    if q in st.session_state.new_questions:
        q["type"] = st.selectbox("Question Type", ["Open", "Closed"], index=0 if q["type"] == "Open" else 1, key=f"type_{i}")

    q["text"] = st.text_area(f"{i+1}. ({q['type']})", value=q["text"], key=f"text_{i}")

    if q["type"] == "Closed":
        remove_indices = []
        for j, opt in enumerate(q["options"]):
            cols = st.columns([5, 1])
            q["options"][j] = cols[0].text_input(f"Option {j+1}", value=opt, key=f"opt_{i}_{j}")
            if cols[1].button("❌", key=f"delopt_{i}_{j}"):
                remove_indices.append(j)
        for index in sorted(remove_indices, reverse=True):
            q["options"].pop(index)
            st.rerun()
        if st.button("➕ Add Option", key=f"addopt_{i}"):
            q["options"].append("")
            st.rerun()
    edited_questions.append((q["type"], q["text"], q["options"]))

st.subheader("Fill Placeholders")
insert_fields = set()
pattern = r"{[iI]nsert (.*?)}"
for qtype, text, opts in edited_questions:
    insert_fields.update(re.findall(pattern, text))
    for opt in opts:
        insert_fields.update(re.findall(pattern, opt))

replacements = {}
for field in sorted(insert_fields):
    replacements[field] = st.text_input(f"{field}:")

if st.button("📤 Export to Word"):
    for question in st.session_state.db_questions:
        q_key = question["id"]
        q_type, q_id_str = q_key.split("_")  
        q_id = int(q_id_str)
        supabase.rpc("increment_print_count", {"q_id": q_id, "q_type": q_type.lower()}).execute()
       
        response = supabase.rpc("increment_print_count", {"q_id": q_id, "q_type": q_type.lower()}).execute()
        st.write("RPC response:", response)
    
    doc = Document()
    doc.add_heading(survey_title, 0)
    doc.add_paragraph(survey_intro)
    doc.add_heading("Questions", level=1)
    for i, (qtype, text, options) in enumerate(edited_questions, 1):
        for k, v in replacements.items():
            text = text.replace(f"{{Insert {k}}}", v).replace(f"{{insert {k}}}", v)
            options = [opt.replace(f"{{Insert {k}}}", v).replace(f"{{insert {k}}}", v) for opt in options]
        doc.add_paragraph(f"{i}. ({qtype}) {text}", style="List Number")
        for opt in options:
            doc.add_paragraph(f"- {opt}", style="List Bullet")
    doc.save("survey.docx")
    with open("survey.docx", "rb") as f:
        st.download_button("📥 Download Survey", f, file_name="survey.docx")
